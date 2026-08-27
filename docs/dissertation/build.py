"""Compile the dissertation to the WMG MSc template: docs/dissertation/*.md -> dissertation.docx.

Windows-first usage:
    python docs/dissertation/build.py            # build docx (+ combined md)
    python docs/dissertation/build.py --md-only  # combined markdown only

Produces the WMG front matter (submission pro-forma, title page, declaration, abstract,
acknowledgements, auto table of contents + lists of tables/figures), then Chapters 1-7,
References and Appendices, with auto-numbered captions and footer page numbers. Uses pandoc if it
is on PATH; otherwise the built-in python-docx renderer IS the production path (pandoc is not
installed on the build machine). The renderer covers exactly what these chapters use: ATX
headings, paragraphs, GFM tables, bullet/number lists, fenced code, images, inline bold/italic/
code. TOC/list fields are inserted as Word fields — after opening the .docx, select all and press
F9 (or right-click -> Update Field) to populate page numbers.
"""
from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
TITLE = "Explainable Machine-Learning Credit Scoring for Green Loans under Extreme Class Imbalance"
COURSE = "Applied Artificial Intelligence"
SUBMITTED = "September 2026"

# Candidate details. These are written into the pro-forma and title page on every build, so they
# must live here rather than being typed into the .docx - the build overwrites the front matter.
STUDENT_NAME = "Tatphong Kruerattanakul"
STUDENT_ID = "5700836"
AUTHOR_LINE = "TATPHONG KRUERATTANAKUL, MSc Applied Artificial Intelligence"
TEMPLATE_PATH = HERE / "24-25_wmg_ft_msc_dissertation_template.docx"
# The WMG mark, extracted from the template's own media (word/media/image1.jpeg) so the
# document carries the same asset the template supplies rather than a look-alike. The template
# places it once, at the head of the title page, and this build does the same.
LOGO_PATH = HERE / "assets" / "wmg_logo.jpeg"
LOGO_WIDTH_IN = 3.86         # the width the template prints it at

# Front matter is built explicitly; the body loop renders chapters 01..99 (00_abstract is folded
# into the front matter, so it is excluded there).
BODY_GLOB = "[0-9][0-9]_*.md"
ABSTRACT_FILE = "00_abstract.md"


def body_files() -> list[Path]:
    return [p for p in sorted(HERE.glob(BODY_GLOB)) if p.name != ABSTRACT_FILE]


def combined_markdown() -> str:
    parts = [f"% {TITLE}\n% MSc {COURSE} — University of Warwick\n"]
    parts.append((HERE / ABSTRACT_FILE).read_text(encoding="utf-8").strip() + "\n")
    for p in body_files():
        parts.append(p.read_text(encoding="utf-8").strip() + "\n")
    return "\n\n".join(parts)


# ------------------------------------------------------------------ pandoc path
def build_with_pandoc(md_path: Path, out_path: Path) -> None:
    subprocess.run(
        ["pandoc", str(md_path), "-o", str(out_path),
         "--from", "markdown+pipe_tables", "--toc", "--toc-depth", "3",
         "--resource-path", str(HERE)],
        check=True,
    )


# ------------------------------------------------------------------ OXML field helpers
def _qn(tag: str):
    from docx.oxml.ns import qn
    return qn(tag)


def _field(paragraph, instruction: str, cached: str = "") -> None:
    """Insert a Word complex field (begin / instrText / separate / cached / end) into a run."""
    from docx.oxml import OxmlElement

    def _mk(tag, **attrs):
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(_qn(k), v)
        return el

    run = paragraph.add_run()
    run._r.append(_mk("w:fldChar", **{"w:fldCharType": "begin"}))
    instr = _mk("w:instrText", **{"xml:space": "preserve"})
    instr.text = instruction
    run._r.append(instr)
    run._r.append(_mk("w:fldChar", **{"w:fldCharType": "separate"}))
    if cached:
        t = _mk("w:t"); t.text = cached
        run._r.append(t)
    run._r.append(_mk("w:fldChar", **{"w:fldCharType": "end"}))


def _footer_page_numbers(section) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(p, "PAGE", cached="1")


def _set_page_numbering(section, fmt: str, start: int | None = None) -> None:
    """Set a section's page-number format (e.g. 'lowerRoman' front matter, 'decimal' body)."""
    from docx.oxml import OxmlElement
    sectPr = section._sectPr
    for old in sectPr.findall(_qn("w:pgNumType")):
        sectPr.remove(old)
    el = OxmlElement("w:pgNumType")
    el.set(_qn("w:fmt"), fmt)
    if start is not None:
        el.set(_qn("w:start"), str(start))
    sectPr.append(el)


def _strip_heading_numbering(doc) -> None:
    """Remove the template's automatic multilevel numbering from the Heading styles.

    The WMG template binds Heading 1-4 to a numbered list, so Word prepends its own "1.1" to a
    heading whose text already reads "1.1 Motivation" (producing "1.1 1.1 Motivation" in the
    document and the contents page, plus a list marker on every heading). The chapter numbers
    are authored in the Markdown source and every cross-reference in the text refers to them, so
    the authored numbers are kept and Word's automatic numbering is removed.
    """
    for level in range(1, 10):
        try:
            st = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        pPr = st.element.find(_qn("w:pPr"))
        if pPr is None:
            continue
        for numPr in pPr.findall(_qn("w:numPr")):
            pPr.remove(numPr)


def _strip_heading_outline(doc) -> None:
    """Remove the outline level from the Heading styles.

    Word draws a collapse/expand triangle ("dropdown") beside every paragraph that carries an
    outline level, and shows the same tree in the Navigation pane. The contents page is
    unaffected because the TOC fields are built with the ``\\t`` switch (collect by style name)
    rather than ``\\o`` (collect by outline level) - see ``_front_matter``.

    Deleting the element is not enough: Word falls back to the built-in definition of
    ``Heading N`` (outline level N-1), so the arrows come back. The level must be explicitly
    overridden to 9, which is Word's value for body text.
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement

    def force_body_level(pPr):
        for ol in pPr.findall(_qn("w:outlineLvl")):
            pPr.remove(ol)
        el = OxmlElement("w:outlineLvl")
        el.set(_qn("w:val"), "9")          # 9 = body text, i.e. no outline level
        pPr.append(el)

    names = [f"Heading {n}" for n in range(1, 10)]
    # The template's front-matter styles are based on Heading 1/2, so they inherit an outline
    # level (and therefore a collapse control) unless it is overridden here as well.
    names += ["Alt Heading 1", "Alt Heading 2", "Title", "TOC Heading"]
    for name in names:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        if st.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        pPr = st.element.find(_qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            st.element.insert(0, pPr)
        force_body_level(pPr)

    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        pPr = p._p.find(_qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._p.insert(0, pPr)
        force_body_level(pPr)


# Custom heading styles. Word enforces its own semantics on the *built-in* Heading 1-4 styles:
# whatever outline level the file specifies, Word restores the built-in one when the document is
# opened and saved, which brings the collapse/expand arrows back. Custom styles based on Normal
# are left alone, so the headings below are rendered with these instead. The TOC fields collect
# by style name, so the contents page is unaffected.
HEADING_STYLES = {
    1: ("Chapter Heading", 16, True),
    2: ("Section Heading", 13, True),
    3: ("Subsection Heading", 12, True),
    4: ("Minor Heading", 11, False),
}


def _make_heading_styles(doc) -> None:
    """Create the custom heading styles (based on Normal, explicitly outline level 9)."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    for level, (name, size, bold) in HEADING_STYLES.items():
        try:
            st = doc.styles[name]
        except KeyError:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles["Normal"]
        st.quick_style = True
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.italic = not bold and level == 4
        st.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)      # the template's heading blue
        pf = st.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(14 if level == 1 else 10)
        pf.space_after = Pt(4)
        pf.keep_with_next = True
        pPr = st.element.get_or_add_pPr()
        for ol in pPr.findall(_qn("w:outlineLvl")):
            pPr.remove(ol)
        el = OxmlElement("w:outlineLvl")
        el.set(_qn("w:val"), "9")                            # body text: no collapse control
        pPr.append(el)


def _add_heading(doc, text: str, level: int):
    """Add a heading using the custom styles, falling back to built-ins if they are missing."""
    name = HEADING_STYLES.get(min(level, 4), (None,))[0]
    if name:
        try:
            return doc.add_paragraph(text, style=name)
        except KeyError:
            pass
    return doc.add_heading(text, level=min(level, 4))


def _enable_update_fields(doc) -> None:
    """Ask Word to update all fields (TOC, lists, page numbers) when the document is opened."""
    from docx.oxml import OxmlElement
    settings = doc.settings.element
    if settings.find(_qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(_qn("w:val"), "true")
        settings.append(el)


def _apply_professional_styles(doc) -> None:
    """Academic house style: A4, 2.54 cm margins, Times New Roman 12/1.5 justified body,
    a clear bold heading hierarchy, and italic centred captions."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Mm, Pt

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)          # A4 portrait
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, side, Cm(2.54))

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size in ((1, 16), (2, 13.5), (3, 12), (4, 11)):
        st = doc.styles[f"Heading {level}"]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = None                                # keep formal black
        hpf = st.paragraph_format
        hpf.space_before = Pt(14 if level == 1 else 10)
        hpf.space_after = Pt(4)
        hpf.keep_with_next = True
        hpf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    try:
        cap = doc.styles["Caption"]
        cap.font.name = "Times New Roman"
        cap.font.size = Pt(10.5)
        cap.font.italic = True
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
    except KeyError:
        pass


# ------------------------------------------------------------------ inline / block renderers
_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`[^`]+`)")


def _add_inline(par, text: str) -> None:
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            par.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            par.add_run(tok[1:-1]).font.name = "Consolas"
        else:
            par.add_run(tok)


def _list_para(doc, raw: str, ordered: bool = False, idx: int = 1):
    """A list item that prefers real list styles (present in from-scratch docs) and falls back
    to an indented paragraph with a literal marker when the template lacks those styles."""
    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
        _add_inline(p, raw)
        return p
    except KeyError:
        pass
    # The WMG template has no List Bullet/Number styles, so the marker is written literally.
    # A hanging indent (not a tab) keeps the wrapped lines aligned under the text, which is why
    # the marker is followed by a single space rather than a tab stop.
    from docx.shared import Cm, Pt
    try:
        p = doc.add_paragraph(style="List Paragraph")
    except KeyError:
        p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-0.5)
    pf.space_after = Pt(3)
    _add_inline(p, (f"{idx}. " if ordered else "• ") + raw)
    return p


def _add_table(doc, rows: list[str]) -> None:
    parsed = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        if set("".join(cells)) <= set("-: "):
            continue
        parsed.append(cells)
    if not parsed:
        return
    ncols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=ncols)
    for style_name in ("Light Grid Accent 1", "Table Grid"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    for i, r in enumerate(parsed):
        for j in range(ncols):
            cell_par = table.rows[i].cells[j].paragraphs[0]
            _add_inline(cell_par, r[j] if j < len(r) else "")
            if i == 0:
                for run in cell_par.runs:
                    run.bold = True


def _caption(doc, kind: str, title: str) -> None:
    """Caption-style paragraph with a SEQ field so the Lists of Tables/Figures populate."""
    try:
        p = doc.add_paragraph(style="Caption")
    except KeyError:
        p = doc.add_paragraph()
    p.add_run(f"{kind} ")
    _field(p, f"SEQ {kind} \\* ARABIC", cached="1")
    _add_inline(p, f": {title}")


# ------------------------------------------------------------------ front matter
# When building on the WMG template, front-matter titles use the template's own named styles
# ("Alt Heading 1", "Title") so they inherit its house look; otherwise a bold paragraph is used.
FRONT_HEADING_STYLE = None   # set to "Alt Heading 1" in template mode
FRONT_TITLE_STYLE = None     # set to "Title" in template mode


def _bold_title(doc, text: str, size: int = 16, center: bool = False, style: str | None = None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    if style:
        try:
            p = doc.add_paragraph(text, style=style)
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return p
        except KeyError:
            pass
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    return p


def _add_logo(doc, width_in: float = LOGO_WIDTH_IN) -> None:
    """Place the WMG mark, centred. Silently skipped if the asset is missing, so a checkout
    without it still builds."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    if not LOGO_PATH.exists():
        print(f"  [warn] logo not found: {LOGO_PATH}", file=sys.stderr)
        return
    doc.add_picture(str(LOGO_PATH), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _front_matter(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    def para(text="", center=False, italic=False, bold=False, size=None):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            r = p.add_run(text); r.italic = italic; r.bold = bold
            if size:
                r.font.size = Pt(size)
        return p

    # 1. Project Submission Pro-Forma - wording and checklist follow the WMG template
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Cm

    def tick(text, checked=True):
        """A checklist line as the template sets it: a box, then the item."""
        pp = doc.add_paragraph()
        pp.paragraph_format.left_indent = Cm(0.6)
        pp.paragraph_format.space_after = Pt(2)
        pp.add_run(("☒ " if checked else "☐ ") + text)
        return pp

    _bold_title(doc, "Project Submission Pro-Forma", 15, style=FRONT_HEADING_STYLE)
    para(f"Student name:  {STUDENT_NAME}")
    para(f"Student ID:  {STUDENT_ID}")
    para()
    para("I wish the dissertation to be considered for the course (select one only):")
    tick(f"MSc in {COURSE}")
    para()
    para("I confirm that I have included in my dissertation:")
    tick("An abstract of the work completed")
    tick("A declaration of my contribution to the work and its suitability for the degree")
    tick("A table of contents")
    tick("A list of figures & tables (if applicable)")
    tick("A glossary of terms (where appropriate)", checked=False)
    tick("A clear statement of my project objectives")
    tick("A full reference list (the Harvard referencing style is recommended for WMG)")
    tick("An appendix containing email confirmation of ethical approval or waiver", checked=False)
    para()

    # Ethical-approval line: the parts the candidate must complete are highlighted, as they are
    # in the template, so nothing is submitted with a placeholder left in by accident.
    pe = doc.add_paragraph()
    pe.add_run("If receiving ethical approval, the ethical approval number for this research is: ")
    pe.add_run("insert reference number by replacing this highlighted text"
               ).font.highlight_color = WD_COLOR_INDEX.YELLOW
    pe.add_run(". If ethical approval was waived, then please write ")
    pe.add_run("‘Ethical Approval Waived’"
               ).font.highlight_color = WD_COLOR_INDEX.YELLOW
    pe.add_run(". Remember to include your email confirmation of ethical approval or waiver in "
               "the Appendix of your submission (Appendix B).")
    para()
    tick("I consent to ongoing storage of this dissertation and potential access by third "
         "parties (e.g. for staff/student training purposes)")
    para()
    para("Signed: …………………………………………….          Date: ………………..")
    doc.add_page_break()

    # 2. Title page
    para()
    _add_logo(doc)
    for _ in range(2):
        para()
    _bold_title(doc, TITLE, 20, center=True, style=FRONT_TITLE_STYLE)
    para()
    para("by", center=True)
    para()
    para(AUTHOR_LINE, center=True)
    para()
    para(f"Dissertation submitted in partial fulfilment for the Degree of Master of Science in "
         f"{COURSE}", center=True)
    para()
    para(f"Submitted {SUBMITTED}", center=True)
    doc.add_page_break()

    # 3. Declaration (verbatim from the WMG template)
    _bold_title(doc, "Declaration", 15, style=FRONT_HEADING_STYLE)
    para("I have read and understood the rules on cheating, plagiarism and appropriate referencing "
         "as outlined in my handbook and I declare that the work contained in this assignment is my "
         "own, unless otherwise acknowledged.")
    para("No substantial part of the work submitted here has also been submitted by me in other "
         "assessments for this or previous degree courses, and I acknowledge that if this has been "
         "done an appropriate reduction in the mark I might otherwise have received will be made.")
    para()
    para("Signed: …………………………………………….          Date: ………………..")
    doc.add_page_break()

    # 4. Abstract (body of 00_abstract.md, minus its own H1)
    _bold_title(doc, "Abstract", 15, style=FRONT_HEADING_STYLE)
    for line in (HERE / ABSTRACT_FILE).read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if not s or s.startswith("# "):
            continue
        _add_inline(doc.add_paragraph(), s)
    doc.add_page_break()

    # 5. Acknowledgements
    _bold_title(doc, "Acknowledgements", 15, style=FRONT_HEADING_STYLE)
    para("[Optional: acknowledgements of supervisor, organisation, and personal support.]")
    doc.add_page_break()

    # 6. Table of contents
    _bold_title(doc, "Table of Contents", 15, style=FRONT_HEADING_STYLE)
    _field(doc.add_paragraph(),
           # \t collects by style name; \u must NOT be present, because it would ask Word to
           # collect by applied outline level, which is deliberately set to body text (9) so
           # that no collapse controls are drawn.
           'TOC \\t "Chapter Heading,1,Section Heading,2,Subsection Heading,3" \\h \\z',
           cached="Right-click and Update Field to populate the table of contents.")
    doc.add_page_break()

    # 7. Lists of tables / figures
    _bold_title(doc, "List of Tables", 15, style=FRONT_HEADING_STYLE)
    _field(doc.add_paragraph(), 'TOC \\h \\z \\c "Table"',
           cached="Right-click and Update Field to populate the list of tables.")
    para()
    _bold_title(doc, "List of Figures", 15, style=FRONT_HEADING_STYLE)
    _field(doc.add_paragraph(), 'TOC \\h \\z \\c "Figure"',
           cached="Right-click and Update Field to populate the list of figures.")
    # No trailing page break: the body section break (build_with_docx) provides the transition
    # and switches page numbering from roman (front matter) to arabic (body).


# ------------------------------------------------------------------ body
_IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def _render_chapter(doc, chapter: Path) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    lines = chapter.read_text(encoding="utf-8").splitlines()
    last_heading = chapter.stem
    i, in_code, code_buf, table_buf = 0, False, [], []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            _caption(doc, "Table", last_heading)   # caption ABOVE the table
            _add_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            flush_table()
            if in_code:
                r = doc.add_paragraph().add_run("\n".join(code_buf))
                r.font.name = "Consolas"; r.font.size = Pt(9)
                code_buf = []
            in_code = not in_code
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        if line.lstrip().startswith("|"):
            table_buf.append(line); i += 1; continue
        flush_table()

        m = _IMG_RE.match(line.strip())
        if m:
            alt, rel = m.group(1), m.group(2)
            img = (chapter.parent / rel).resolve()
            if img.exists():
                doc.add_picture(str(img), width=Inches(5.8))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _caption(doc, "Figure", alt)       # caption BELOW the image
            else:
                print(f"  [warn] missing figure: {img}", file=sys.stderr)
            i += 1; continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("# ").strip()
            last_heading = text
            _add_heading(doc, text, level)
            i += 1; continue

        stripped = line.strip()
        if re.match(r"^[-*]\s+", stripped):
            _list_para(doc, re.sub(r"^[-*]\s+", "", stripped), ordered=False)
            i += 1; continue
        mnum = re.match(r"^(\d+)\.\s+", stripped)
        if mnum:
            _list_para(doc, re.sub(r"^\d+\.\s+", "", stripped), ordered=True, idx=int(mnum.group(1)))
            i += 1; continue

        if stripped:
            buf = [stripped]
            while (i + 1 < len(lines) and lines[i + 1].strip()
                   and not lines[i + 1].lstrip().startswith(("#", "|", "-", "*", "```", "!["))
                   and not re.match(r"^\d+\.\s+", lines[i + 1].strip())):
                i += 1
                buf.append(lines[i].strip())
            _add_inline(doc.add_paragraph(), " ".join(buf))
        i += 1

    flush_table()
    doc.add_page_break()


def _assemble(doc) -> None:
    """Front matter (roman page numbers) + a body section (arabic) + all chapters."""
    from docx.enum.section import WD_SECTION

    _footer_page_numbers(doc.sections[0])
    _set_page_numbering(doc.sections[0], "lowerRoman", start=1)
    _front_matter(doc)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.footer.is_linked_to_previous = False
    _footer_page_numbers(body)
    _set_page_numbering(body, "decimal", start=1)

    for chapter in body_files():
        _render_chapter(doc, chapter)

    _enable_update_fields(doc)


def build_with_docx(out_path: Path) -> None:
    """From-scratch build with our own academic house style (used if the template is absent)."""
    global FRONT_HEADING_STYLE, FRONT_TITLE_STYLE
    from docx import Document

    FRONT_HEADING_STYLE = FRONT_TITLE_STYLE = None
    doc = Document()
    _apply_professional_styles(doc)
    _make_heading_styles(doc)
    _assemble(doc)
    doc.save(out_path)


def _clear_body(doc) -> None:
    """Remove ALL of the template's placeholder content — paragraphs, tables, and structured
    document tags (Word wraps its TOC in a <w:sdt>, which is neither w:p nor w:tbl) — keeping
    only the trailing sectPr (page setup + header/footer references) so the house theme, styles
    and margins survive."""
    body = doc.element.body
    keep = _qn("w:sectPr")
    for child in list(body):
        if child.tag != keep:
            body.remove(child)


def build_from_template(out_path: Path) -> None:
    """Populate the actual WMG template .docx: inherit its styles, theme, fonts, margins and
    page-numbering, and pour our content in using the template's own named styles."""
    global FRONT_HEADING_STYLE, FRONT_TITLE_STYLE
    from docx import Document

    doc = Document(str(TEMPLATE_PATH))
    _clear_body(doc)
    _strip_heading_numbering(doc)
    _strip_heading_outline(doc)                 # style definitions, before content exists
    _make_heading_styles(doc)                   # custom, non-built-in heading styles
    FRONT_HEADING_STYLE, FRONT_TITLE_STYLE = "Alt Heading 1", "Title"
    _assemble(doc)
    _strip_heading_outline(doc)                 # again, to stamp the heading paragraphs
    doc.save(out_path)


# ------------------------------------------------------------------ main
def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(prog="python docs/dissertation/build.py")
    ap.add_argument("--md-only", action="store_true", help="write combined markdown only")
    ap.add_argument("--scratch", action="store_true",
                    help="ignore the WMG template and build with our own house style")
    args = ap.parse_args(argv)

    md_out = HERE / "dissertation.md"
    md_out.write_text(combined_markdown(), encoding="utf-8")
    print(f"[build] combined markdown -> {md_out}")
    if args.md_only:
        return 0

    docx_out = HERE / "dissertation.docx"
    if shutil.which("pandoc"):
        build_with_pandoc(md_out, docx_out)
        print(f"[build] pandoc -> {docx_out}")
    elif TEMPLATE_PATH.exists() and not args.scratch:
        build_from_template(docx_out)
        print(f"[build] populated the WMG template -> {docx_out}")
        print("[build] NOTE: Word will offer to update fields on open (updateFields=true); "
              "accept it to fill the contents/lists/page numbers, or select all and press F9.")
    else:
        build_with_docx(docx_out)
        print(f"[build] python-docx (house style, no template) -> {docx_out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
